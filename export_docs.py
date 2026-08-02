from __future__ import annotations

import argparse
import base64
import hashlib
import html
import json
import os
import re
import sys
from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Optional
from urllib.error import URLError
from urllib.request import Request, urlopen

_BUNDLED_PYTHON = Path(r"C:\Users\worth\.cache\codex-runtimes\codex-primary-runtime\dependencies\python\python.exe")
if sys.version_info[:2] != (3, 12) and _BUNDLED_PYTHON.exists():
    os.execv(str(_BUNDLED_PYTHON), [str(_BUNDLED_PYTHON), str(Path(__file__).resolve()), *sys.argv[1:]])

from lxml import html as lxml_html
from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from import_mbox import (
    decode_html_entities,
    sanitize_filename,
)


MONTHS = {
    "january": 1,
    "february": 2,
    "march": 3,
    "april": 4,
    "may": 5,
    "june": 6,
    "july": 7,
    "august": 8,
    "september": 9,
    "october": 10,
    "november": 11,
    "december": 12,
}


TYPE_ORDER = ("backThen", "retrospect")
TYPE_TO_FOLDER = {
    "backThen": "Back Then History",
    "retrospect": "The Retrospect",
}


@dataclass
class JsonMessage:
    source_file: str
    message_id: str
    sender: str
    kind: str
    folder_name: str
    subject: str
    received_date: datetime
    historical_date: Optional[datetime]
    event_title: str
    body_text: str
    body_html: str
    raw_source: str
    headers: list[dict[str, str]]
    mime: dict[str, Any]


def parse_historical_date(text: str) -> Optional[datetime]:
    patterns = (
        re.compile(r"On\s+([A-Z][a-z]+)\s+(\d{1,2}),\s+(\d{4})", re.I),
        re.compile(r"\b([A-Z][a-z]+)\s+(\d{1,2}),\s+(\d{4})\b"),
    )
    for pattern in patterns:
        match = pattern.search(str(text or ""))
        if not match:
            continue
        month = MONTHS.get(match.group(1).lower(), 0)
        day = int(match.group(2))
        year = int(match.group(3))
        if month and day and year:
            try:
                return datetime(year, month, day)
            except ValueError:
                return None
    return None


def parse_date(value: str) -> datetime:
    parsed = datetime.fromisoformat(value.replace("Z", "+00:00"))
    return parsed


def stable_id(message: JsonMessage) -> str:
    source = "\n".join(
        [
            message.message_id,
            message.sender,
            message.subject,
            message.received_date.isoformat(),
            message.event_title,
        ]
    )
    return hashlib.sha1(source.encode("utf-8")).hexdigest()[:10]


def load_messages(json_path: Path) -> list[JsonMessage]:
    payload = json.loads(json_path.read_text(encoding="utf-8"))
    messages: list[JsonMessage] = []
    for item in payload.get("messages", []):
        historical_date = item.get("historical_date")
        messages.append(
            JsonMessage(
                source_file=str(item.get("source_file", "")),
                message_id=str(item.get("message_id", "") or ""),
                sender=str(item.get("sender", "") or ""),
                kind=str(item.get("kind", "") or ""),
                folder_name=str(item.get("folder_name", "") or ""),
                subject=str(item.get("subject", "") or ""),
                received_date=parse_date(str(item.get("received_date", ""))),
                historical_date=parse_date(historical_date) if historical_date else None,
                event_title=str(item.get("event_title", "") or ""),
                body_text=str(item.get("body_text", "") or ""),
                body_html=str(item.get("body_html", "") or ""),
                raw_source=str(item.get("raw_source", "") or ""),
                headers=list(item.get("headers", [])),
                mime=dict(item.get("mime", {})),
            )
        )
    return messages


def set_run_font(run, family: str, size: int, bold: bool = False, color: str = "000000") -> None:
    run.font.name = family
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), family)


def set_paragraph_format(paragraph, before: int = 0, after: int = 0, line: Optional[float] = None) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if line is not None:
        fmt.line_spacing = line


def clear_cell(cell) -> None:
    tc = cell._tc
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.replace("#", ""))


def set_cell_margins(cell, top: int, left: int, bottom: int, right: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders_none(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def add_horizontal_rule(document: Document, color: str = "CCCECE") -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(paragraph, before=0, after=0)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def compress_text(text: str) -> str:
    output = str(text or "").replace("\r", "")
    output = re.sub(r"[ \t]+\n", "\n", output)
    output = re.sub(r"\n{3,}", "\n\n", output)
    output = re.sub(r"Follow us on our social channels!?\s*", "", output, flags=re.I)
    output = re.sub(r"Unsubscribe[\s\S]*$", "", output, flags=re.I)
    output = output.strip()
    if len(output) > 4000:
        output = output[:4000]
        last_period = output.rfind(". ")
        if last_period > 1000:
            output = output[: last_period + 1]
        output = output.strip() + "\n\n[Content truncated to fit a single page.]"
    return output


def normalize_paragraphs(text: str) -> list[str]:
    return [
        paragraph.strip()
        for paragraph in re.split(r"\n{2,}", str(text or "").replace("\r", ""))
        if paragraph.strip()
    ]


def html_to_text(body_html: str) -> str:
    try:
        tree = lxml_html.fromstring(body_html or "<html></html>")
        return tree.text_content()
    except Exception:
        return re.sub(r"<[^>]+>", "", body_html or "")


def decode_html_entities_text(text: str) -> str:
    return decode_html_entities(text)


def find_html_token_index(html_source: str, token: str) -> int:
    return str(html_source or "").lower().find(str(token or "").lower())


def extract_image_urls(html_source: str) -> list[str]:
    urls: list[str] = []
    seen: set[str] = set()
    pattern = re.compile(r'<img\b[^>]*\bsrc\s*=\s*["\']([^"\']+)["\'][^>]*>', re.I)
    for match in pattern.finditer(str(html_source or "")):
        src = decode_html_entities(str(match.group(1) or "").strip())
        if not src or src.lower().startswith("data:"):
            continue
        if re.search(r"ssl\.gstatic\.com/ui/v1/icons/mail|googleusercontent\.com/a/default-user|cleardot\.gif|logo_loading_|logo_gmail_lockup", src, re.I):
            continue
        if src not in seen:
            seen.add(src)
            urls.append(src)
    return urls


def is_gmail_ui_image_url(url: str) -> bool:
    return bool(re.search(r"ssl\.gstatic\.com/ui/v1/icons/mail|googleusercontent\.com/a/default-user|cleardot\.gif|logo_loading_|logo_gmail_lockup", str(url or ""), re.I))


def is_likely_social_icon_url(url: str) -> bool:
    value = str(url or "").lower()
    return any(token in value for token in ("instagram", "facebook", "linkedin", "twitter", "x-logo", "social", "share", "youtube", "pinterest", "tiktok", "threads", "icon"))


def collect_image_assets(node: Any, assets: list[dict[str, Any]]) -> None:
    if isinstance(node, dict):
        if str(node.get("content_type", "")).lower().startswith("image/") and node.get("payload_base64"):
            try:
                payload = base64.b64decode(str(node["payload_base64"]))
            except Exception:
                payload = b""
            if payload:
                assets.append(
                    {
                        "content_id": node.get("content_id"),
                        "filename": node.get("filename"),
                        "content_type": node.get("content_type"),
                        "bytes": payload,
                        "size": len(payload),
                    }
                )
        for child in node.get("parts", []) or []:
            collect_image_assets(child, assets)


def fetch_url_bytes(url: str) -> Optional[bytes]:
    request = Request(url, headers={"User-Agent": "Mozilla/5.0"})
    try:
        with urlopen(request, timeout=15) as response:
            if 200 <= getattr(response, "status", 200) < 300:
                data = response.read()
                if is_valid_image_bytes(data):
                    return data
    except (URLError, TimeoutError, ValueError, OSError):
        return None
    return None


def image_bytes_to_dimensions(data: bytes) -> tuple[int, int]:
    with Image.open(BytesIO(data)) as image:
        return image.size


def is_valid_image_bytes(data: bytes) -> bool:
    try:
        with Image.open(BytesIO(data)) as image:
            image.verify()
        return True
    except Exception:
        return False


def normalize_image_bytes_for_docx(data: bytes) -> Optional[bytes]:
    try:
        with Image.open(BytesIO(data)) as image:
            image.load()
            converted = image.convert("RGBA") if image.mode not in ("RGB", "RGBA") else image.copy()
            output = BytesIO()
            converted.save(output, format="PNG")
            return output.getvalue()
    except Exception:
        return None


def add_picture_from_bytes(paragraph, data: bytes, width_inches: Optional[float] = None) -> None:
    run = paragraph.add_run()
    if width_inches is None:
        run.add_picture(BytesIO(data))
    else:
        run.add_picture(BytesIO(data), width=Inches(width_inches))


def add_centered_image(document: Document, data: bytes, width_inches: float, space_after: int = 0) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(paragraph, before=0, after=space_after)
    add_picture_from_bytes(paragraph, data, width_inches)


def add_centered_image_with_max_width(document: Document, data: bytes, max_width_inches: float, space_after: int = 0) -> None:
    image_width_px, _ = image_bytes_to_dimensions(data)
    natural_width_inches = image_width_px / 96.0
    width_inches = min(natural_width_inches, max_width_inches)
    add_centered_image(document, data, width_inches, space_after=space_after)


def add_floating_image_to_paragraph(paragraph, data: bytes, width_inches: float, float_side: str = "left") -> None:
    run = paragraph.add_run()
    inline_shape = run.add_picture(BytesIO(data), width=Inches(width_inches))
    inline = inline_shape._inline

    anchor = OxmlElement("wp:anchor")
    side_gap = "137160"
    dist_left = side_gap if float_side == "right" else "0"
    dist_right = side_gap if float_side != "right" else "0"
    for attr, value in (
        ("distT", "0"),
        ("distB", "0"),
        ("distL", dist_left),
        ("distR", dist_right),
        ("simplePos", "0"),
        ("relativeHeight", "0"),
        ("behindDoc", "0"),
        ("locked", "0"),
        ("layoutInCell", "1"),
        ("allowOverlap", "1"),
    ):
        anchor.set(attr, value)

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")

    position_h = OxmlElement("wp:positionH")
    position_h.set("relativeFrom", "margin")
    position_h_offset = OxmlElement("wp:posOffset")
    position_h_offset.text = "0"
    position_h.append(position_h_offset)

    position_v = OxmlElement("wp:positionV")
    position_v.set("relativeFrom", "paragraph")
    position_v_offset = OxmlElement("wp:posOffset")
    position_v_offset.text = "9144"
    position_v.append(position_v_offset)

    extent = deepcopy(inline.find(qn("wp:extent")))
    effect_extent = OxmlElement("wp:effectExtent")
    for side in ("l", "t", "r", "b"):
        effect_extent.set(side, "0")
    wrap_square = OxmlElement("wp:wrapSquare")
    wrap_square.set("wrapText", "bothSides")
    doc_pr = deepcopy(inline.find(qn("wp:docPr")))
    c_nv = deepcopy(inline.find(qn("wp:cNvGraphicFramePr")))
    graphic = deepcopy(inline.find(qn("a:graphic")))

    for element in (simple_pos, position_h, position_v, extent, effect_extent, wrap_square, doc_pr, c_nv, graphic):
        anchor.append(element)

    inline.getparent().replace(inline, anchor)


def add_floating_image_paragraph(document: Document, data: bytes, width_inches: float, float_side: str = "left") -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(paragraph, before=0, after=0)
    add_floating_image_to_paragraph(paragraph, data, width_inches, float_side=float_side)


def add_text_paragraphs_to(container, paragraphs: list[str], family: str, size: int, color: str = "111827") -> None:
    for paragraph_text in paragraphs:
        paragraph = container.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_format(paragraph, before=0, after=6, line=1.1)
        run = paragraph.add_run(paragraph_text)
        set_run_font(run, family, size, bold=False, color=color)


def pick_retrospect_logo_url(record: JsonMessage) -> Optional[str]:
    html_source = record.body_html
    article_start = find_html_token_index(html_source, "On This Day in History")
    segment = html_source[:article_start] if article_start >= 0 else html_source
    urls = extract_image_urls(segment)
    logo_urls = [url for url in urls if not is_likely_social_icon_url(url) and not is_gmail_ui_image_url(url)]
    return logo_urls[0] if logo_urls else (urls[0] if urls else None)


def pick_retrospect_main_image_urls(record: JsonMessage) -> list[str]:
    urls = [url for url in extract_image_urls(record.body_html) if not is_gmail_ui_image_url(url) and not is_likely_social_icon_url(url)]
    article_urls = [url for url in urls if "/Logo.png" not in url and "logo.png" not in url.lower()]
    return article_urls if article_urls else urls


def pick_back_then_logo_url(record: JsonMessage) -> Optional[str]:
    html_source = record.body_html
    article_start = find_html_token_index(html_source, "The history of")
    segment = html_source[:article_start] if article_start >= 0 else html_source
    urls = extract_image_urls(segment)
    for url in urls:
        if re.search(r"/Logo\.png(\?|$|#)|logo\.png", url, re.I):
            return url
    for url in urls:
        if not is_likely_social_icon_url(url) and not is_gmail_ui_image_url(url):
            return url
    return None


def pick_back_then_article_image_url(record: JsonMessage) -> Optional[str]:
    html_source = record.body_html
    footer_start = find_html_token_index(html_source, "Continue Reading")
    article_start = find_html_token_index(html_source, "Today's Object")
    social_start = find_html_token_index(html_source, "facebook")
    social_instagram_start = find_html_token_index(html_source, "instagram")
    start_index = max(article_start, social_start, social_instagram_start, 0)
    end_index = footer_start if footer_start > start_index else len(html_source)
    segment = html_source[start_index:end_index]
    urls = extract_image_urls(segment)
    filtered = [
        url for url in urls
        if not re.search(r"/Logo\.png(\?|$|#)|logo\.png|cleardot\.gif|logo_gmail_lockup|default-user", url, re.I)
        and not is_likely_social_icon_url(url)
        and not is_gmail_ui_image_url(url)
    ]
    return filtered[-1] if filtered else None


def choose_best_image_bytes(url: Optional[str]) -> Optional[bytes]:
    if not url:
        return None
    data = fetch_url_bytes(url)
    if not data:
        return None
    return normalize_image_bytes_for_docx(data)


def require_image_bytes(url: Optional[str], description: str, record: JsonMessage) -> bytes:
    data = choose_best_image_bytes(url)
    if data is None:
        raise ValueError(
            f"Missing required {description} image "
            f"(subject={record.subject!r}, event_title={record.event_title!r})"
        )
    return data


def add_table_row_text(cell, text: str, family: str, size: int, bold: bool = False, color: str = "000000", align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    set_paragraph_format(paragraph, before=0, after=0)
    run = paragraph.add_run(text)
    set_run_font(run, family, size, bold=bold, color=color)


def add_retrospect_date_bar(document: Document, date_value: datetime, available_width_inches: float) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders_none(table)
    table.columns[0].width = Inches(available_width_inches)
    cell = table.cell(0, 0)
    cell.width = Inches(available_width_inches)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    set_cell_shading(cell, "#f1cd73")
    set_cell_margins(cell, top=0, left=240, bottom=0, right=240)
    if hasattr(date_value, "strftime"):
        date_text = f"{date_value.strftime('%A')}, {date_value.strftime('%B')} {date_value.day}, {date_value.year}".upper()
    else:
        date_text = " "
    add_table_row_text(cell, date_text, "Arial", 10, bold=True, color="1f1b16", align=WD_ALIGN_PARAGRAPH.RIGHT)


def add_back_then_banner(document: Document, record: JsonMessage, available_width_inches: float) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders_none(table)
    table.columns[0].width = Inches(available_width_inches)
    cell = table.cell(0, 0)
    cell.width = Inches(available_width_inches)
    set_cell_shading(cell, "#000000")
    set_cell_margins(cell, top=120, left=200, bottom=120, right=200)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(paragraph, before=0, after=0)

    logo_url = pick_back_then_logo_url(record)
    logo_bytes = require_image_bytes(logo_url, "Back Then History logo", record)
    add_picture_from_bytes(paragraph, logo_bytes, 190 / 96.0)


def body_paragraphs_for_retrospect(record: JsonMessage) -> list[str]:
    paragraphs = normalize_paragraphs(record.body_text)
    return paragraphs[1:] if len(paragraphs) > 1 else []


def body_paragraphs_for_back_then(record: JsonMessage) -> list[str]:
    paragraphs = normalize_paragraphs(record.body_text)
    return paragraphs[1:] if len(paragraphs) > 1 else []


def add_text_paragraphs(document: Document, paragraphs: list[str], family: str, size: int, color: str = "111827") -> None:
    add_text_paragraphs_to(document, paragraphs, family, size, color=color)


def add_article_paragraphs(
    document: Document,
    paragraphs: list[str],
    family: str,
    size: int,
    color: str = "111827",
    image_bytes: Optional[bytes] = None,
    image_width_inches: float = 3.0,
) -> None:
    if not paragraphs:
        return
    first_paragraph, *rest = paragraphs
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(paragraph, before=0, after=6, line=1.1)
    if image_bytes:
        add_floating_image_to_paragraph(paragraph, image_bytes, image_width_inches)
    run = paragraph.add_run(first_paragraph)
    set_run_font(run, family, size, bold=False, color=color)
    if rest:
        add_text_paragraphs_to(document, rest, family, size, color=color)


def render_retrospect(document: Document, record: JsonMessage, available_width_inches: float) -> None:
    banner_date = record.historical_date or record.received_date
    add_retrospect_date_bar(document, banner_date, available_width_inches)

    logo_url = pick_retrospect_logo_url(record)
    logo_bytes = require_image_bytes(logo_url, "Retrospect logo", record)
    add_centered_image(document, logo_bytes, 220 / 96.0, space_after=0)

    add_horizontal_rule(document)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(kicker, before=8, after=4)
    run = kicker.add_run("On This Day in History")
    set_run_font(run, "Georgia", 13, bold=True, color="1f1b16")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(title, before=0, after=8)
    run = title.add_run(record.event_title or record.subject or "Untitled")
    set_run_font(run, "Arial", 22, bold=True, color="6f93ad")

    main_urls = pick_retrospect_main_image_urls(record)
    image_bytes = require_image_bytes(
        main_urls[0] if main_urls else None,
        "Retrospect article",
        record,
    )

    article_paragraphs = body_paragraphs_for_retrospect(record)
    compressed = compress_text("\n\n".join(article_paragraphs))
    add_article_paragraphs(document, normalize_paragraphs(compressed), "Arial", 10, image_bytes=image_bytes, image_width_inches=400 / 96.0)


def render_back_then(document: Document, record: JsonMessage, available_width_inches: float) -> None:
    add_back_then_banner(document, record, available_width_inches)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(kicker, before=14, after=0)
    run = kicker.add_run("THE HISTORY OF")
    set_run_font(run, "Arial", 11, bold=True, color="111827")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(title, before=0, after=8)
    run = title.add_run(record.event_title or record.subject or "Untitled")
    set_run_font(run, "Arial", 22, bold=True, color="ef3b3a")

    article_url = pick_back_then_article_image_url(record)
    image_bytes = require_image_bytes(
        article_url,
        "Back Then History article",
        record,
    )

    article_paragraphs = body_paragraphs_for_back_then(record)
    compressed = compress_text("\n\n".join(article_paragraphs))
    add_article_paragraphs(document, normalize_paragraphs(compressed), "Arial", 11, image_bytes=image_bytes, image_width_inches=400 / 96.0)


def compress_text(text: str) -> str:
    output = str(text or "").replace("\r", "")
    output = re.sub(r"[ \t]+\n", "\n", output)
    output = re.sub(r"\n{3,}", "\n\n", output)
    output = re.sub(r"Follow us on our social channels!?\s*", "", output, flags=re.I)
    output = re.sub(r"Unsubscribe[\s\S]*$", "", output, flags=re.I)
    output = output.strip()
    if len(output) > 4000:
        output = output[:4000]
        last_period = output.rfind(". ")
        if last_period > 1000:
            output = output[: last_period + 1]
        output = output.strip() + "\n\n[Content truncated to fit a single page.]"
    return output


def configure_document(document: Document) -> float:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Pt(28)
    section.bottom_margin = Pt(28)
    section.left_margin = Pt(30)
    section.right_margin = Pt(30)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    return section.page_width.inches - section.left_margin.inches - section.right_margin.inches


def build_document(record: JsonMessage) -> Document:
    document = Document()
    available_width = configure_document(document)
    if record.kind == "backThen":
        render_back_then(document, record, available_width)
    else:
        render_retrospect(document, record, available_width)
    return document


def output_path_for(record: JsonMessage, output_dir: Path) -> Path:
    output_folder = output_dir / record.folder_name
    output_folder.mkdir(parents=True, exist_ok=True)
    date_value = record.historical_date or record.received_date
    date_part = date_value.strftime("%Y-%m-%d")
    title_part = sanitize_filename(record.event_title)
    base_name = f"{date_part}-{title_part}"
    return output_folder / f"{base_name}.docx"


def write_document(record: JsonMessage, output_dir: Path) -> Path:
    document = build_document(record)
    path = output_path_for(record, output_dir)
    document.save(str(path))
    return path


def write_failed_log(output_dir: Path, failures: list[str]) -> Path:
    path = output_dir / "failed.log"
    if failures:
        body = "\n".join(failures) + "\n"
    else:
        body = "No failures.\n"
    path.write_text(body, encoding="utf-8")
    return path


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate DOCX files from the grouped newsletter JSON export.")
    parser.add_argument(
        "--input-dir",
        default="output",
        help="Folder containing Back Then History.json and The Retrospect.json.",
    )
    parser.add_argument(
        "--backthen-limit",
        type=int,
        default=None,
        help="Optional cap for Back Then History docs. Omit to generate all.",
    )
    parser.add_argument(
        "--retrospect-limit",
        type=int,
        default=None,
        help="Optional cap for The Retrospect docs. Omit to generate all.",
    )
    parser.add_argument(
        "--skip-existing",
        action="store_true",
        help="Skip DOCX files that already exist in the output folder.",
    )
    args = parser.parse_args()

    input_dir = Path(args.input_dir)
    output_dir = Path("output")
    output_dir.mkdir(parents=True, exist_ok=True)

    total = 0
    failures: list[str] = []
    for kind in TYPE_ORDER:
        json_path = input_dir / f"{TYPE_TO_FOLDER[kind]}.json"
        if not json_path.exists():
            continue
        messages = load_messages(json_path)
        kind_limit = args.backthen_limit if kind == "backThen" else args.retrospect_limit
        if kind_limit is not None:
            if kind_limit < 0:
                raise SystemExit(f"Invalid limit for {kind}: {kind_limit}")
            messages = messages[:kind_limit]
        for message in messages:
            path = output_path_for(message, output_dir)
            if args.skip_existing and path.exists():
                print(f"Already exists {path}")
                continue
            try:
                path = write_document(message, output_dir)
            except Exception as exc:
                failures.append(f"{path} | {message.subject} | {exc}")
                print(f"FAILED {path}: {exc}")
                continue
            total += 1
            print(path)
    failed_log = write_failed_log(output_dir, failures)
    print(f"Wrote {total} documents to {output_dir}")
    print(f"Wrote failure log to {failed_log}")
    return 1 if failures else 0


if __name__ == "__main__":
    raise SystemExit(main())
